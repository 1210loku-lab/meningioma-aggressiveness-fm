#!/usr/bin/env python3
"""Build a Neuro-Oncology Advances upload package from the docs/ source."""
from pathlib import Path
import re
import shutil

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "manuscript_draft_v0.md"
LEGENDS = ROOT / "docs" / "figure_legends.md"
COVER = ROOT / "docs" / "cover_letter_Neuro_Oncology_Advances.md"
OUT = ROOT / "docs" / "NOA_submission"

FIGURES = [
    ROOT / "results/figures_pub/Figure1_complete.png",
    ROOT / "results/scrna/fig_umap_celltype_v2.png",
    ROOT / "results/scrna/fig_GSE206647_grade_program.png",
    ROOT / "results/figures_pub/fig_classical_drivers.png",
    ROOT / "results/figures_pub/fig_geneformer_embedding.png",
]


def clean_inline(text):
    return text.replace("`", "").strip()


def add_page_number(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Page ")
    run = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def enable_line_numbers(section):
    node = OxmlElement("w:lnNumType")
    node.set(qn("w:countBy"), "1")
    node.set(qn("w:start"), "1")
    node.set(qn("w:restart"), "newPage")
    section._sectPr.append(node)


def setup(doc, line_numbers=True):
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(1)
    sec.left_margin = sec.right_margin = Inches(1)
    if line_numbers:
        enable_line_numbers(sec)
    add_page_number(sec)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2
    normal.paragraph_format.space_after = Pt(0)
    for name, size in (("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 12)):
        st = doc.styles[name]
        st.font.name = "Times New Roman"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(6)


def add_plain_with_citations(paragraph, text, superscript=True):
    citation = re.compile(r"\[(\d+(?:\s*[,;]\s*\d+|\s*[-–]\s*\d+)*)\]")
    pos = 0
    for m in citation.finditer(text):
        paragraph.add_run(text[pos:m.start()])
        run = paragraph.add_run(m.group(1).replace(" ", ""))
        run.font.superscript = superscript
        pos = m.end()
    paragraph.add_run(text[pos:])


def add_inline(paragraph, text, superscript=True):
    text = clean_inline(text)
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            before = len(paragraph.runs)
            add_plain_with_citations(paragraph, part[2:-2], superscript)
            for run in paragraph.runs[before:]:
                run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            before = len(paragraph.runs)
            add_plain_with_citations(paragraph, part[1:-1], superscript)
            for run in paragraph.runs[before:]:
                run.italic = True
        else:
            add_plain_with_citations(paragraph, part, superscript)


def add_table(doc, rows):
    parsed = [[clean_inline(c).replace("**", "") for c in r.strip().strip("|").split("|")] for r in rows]
    if len(parsed) < 3:
        return
    header, body = parsed[0], parsed[2:]
    landscape = len(header) >= 8
    already_landscape = doc.sections[-1].orientation == WD_ORIENT.LANDSCAPE
    if landscape and not already_landscape:
        sec = doc.add_section(WD_SECTION.NEW_PAGE)
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width, sec.page_height = Inches(11), Inches(8.5)
        sec.top_margin = sec.bottom_margin = Inches(0.55)
        sec.left_margin = sec.right_margin = Inches(0.55)
        enable_line_numbers(sec)
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for i, value in enumerate(header):
        table.rows[0].cells[i].text = value
        header_p = table.rows[0].cells[i].paragraphs[0]
        header_p.paragraph_format.line_spacing = 1
        header_p.paragraph_format.space_after = Pt(0)
        for run in header_p.runs:
            run.bold = True
            run.font.size = Pt(7 if landscape else 8)
    for values in body:
        if len(values) != len(header):
            continue
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value
            for p in cells[i].paragraphs:
                p.paragraph_format.line_spacing = 1
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(7 if landscape else 8)
    doc.add_paragraph()
    if landscape:
        sec = doc.add_section(WD_SECTION.NEW_PAGE)
        sec.orientation = WD_ORIENT.PORTRAIT
        sec.page_width, sec.page_height = Inches(8.5), Inches(11)
        sec.top_margin = sec.bottom_margin = Inches(1)
        sec.left_margin = sec.right_margin = Inches(1)
        enable_line_numbers(sec)


def add_markdown(doc, text, strip_doi=False, superscript=True):
    table = []
    first_h1 = True
    for raw in text.splitlines():
        line = raw.rstrip()
        if strip_doi and re.match(r"^\d+\. ", line):
            line = re.sub(r"\s+doi:\S+\s*$", "", line)
        if line.startswith("|"):
            table.append(line)
            continue
        if table:
            add_table(doc, table)
            table = []
        if not line or line == "---":
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            if first_h1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                first_h1 = False
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line[2:].replace("**", ""))
            run.bold = True
            run.font.size = Pt(16)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:], superscript)
        else:
            p = doc.add_paragraph()
            add_inline(p, line, superscript)
            # Figure legends are one paragraph each. Keep a legend intact when it
            # fits on a page so Word does not leave a short orphaned tail.
            if re.match(r"^\*\*Figure\s+\d+\.", line):
                p.paragraph_format.keep_together = True
    if table:
        add_table(doc, table)


def build_main(source):
    supp = source.index("## Supplementary")
    refs = source.index("## References")
    main_text = source[:supp].rstrip() + "\n\n" + source[refs:].strip()
    doc = Document()
    setup(doc)
    add_markdown(doc, main_text, strip_doi=True, superscript=True)
    doc.add_page_break()
    add_markdown(doc, LEGENDS.read_text(), superscript=False)
    path = OUT / "01_Main_Manuscript_NOA.docx"
    doc.save(path)
    return path


def build_supp(source):
    supp = source.index("## Supplementary")
    refs = source.index("## References")
    text = "# Supplementary Material\n\n" + source[supp:refs].replace("## Supplementary", "## Supplementary Methods and Tables")
    text += "\n\n" + source[refs:]
    doc = Document()
    setup(doc)
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Inches(11), Inches(8.5)
    sec.top_margin = sec.bottom_margin = Inches(0.55)
    sec.left_margin = sec.right_margin = Inches(0.55)
    add_markdown(doc, text, strip_doi=True, superscript=True)
    path = OUT / "02_Supplementary_Material_NOA.docx"
    doc.save(path)
    return path


def build_cover():
    doc = Document()
    setup(doc, line_numbers=False)
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(0.7)
    normal = doc.styles["Normal"]
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(5)
    add_markdown(doc, COVER.read_text(), superscript=False)
    path = OUT / "03_Cover_Letter_NOA.docx"
    doc.save(path)
    return path


def export_figures():
    figdir = OUT / "Figures"
    figdir.mkdir(parents=True, exist_ok=True)
    for idx, src in enumerate(FIGURES, 1):
        with Image.open(src) as im:
            if im.mode not in ("RGB", "L"):
                im = im.convert("RGB")
            im.save(figdir / f"Figure_{idx}.tif", compression="tiff_lzw", dpi=(300, 300))
    ga = ROOT / "docs/GA/GA_final.png"
    with Image.open(ga) as im:
        if im.mode not in ("RGB", "L"):
            im = im.convert("RGB")
        im.save(figdir / "Graphical_Abstract.tif", compression="tiff_lzw", dpi=(300, 300))


def copy_tables():
    tabdir = OUT / "Supplementary_Tables"
    tabdir.mkdir(parents=True, exist_ok=True)
    sources = [
        ROOT / "docs/Table_S2_aggressiveness_program_genes.csv",
        ROOT / "docs/Table_S3_GSE206647_pseudobulk_scores.csv",
        ROOT / "docs/Table_S4_translational_drug_candidates.csv",
        ROOT / "results/drug_repurposing/target_evidence_matrix.csv",
    ]
    names = [
        "Table_S2_aggressiveness_program_genes.csv",
        "Table_S3_GSE206647_pseudobulk_scores.csv",
        "Table_S4_translational_drug_candidates.csv",
        "Table_S5_target_biomarker_evidence_matrix.csv",
    ]
    for src, name in zip(sources, names):
        shutil.copy2(src, tabdir / name)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    source = SOURCE.read_text()
    paths = [build_main(source), build_supp(source), build_cover()]
    export_figures()
    copy_tables()
    shutil.copy2(ROOT / "docs/NOA_SUBMISSION_CHECKLIST.md", OUT / "00_README_CHECKLIST.md")
    print("\n".join(str(p) for p in paths))
    print(OUT / "Figures")
    print(OUT / "Supplementary_Tables")


if __name__ == "__main__":
    main()
