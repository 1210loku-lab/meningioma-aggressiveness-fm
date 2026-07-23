from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = ROOT / "docs" / "manuscript_draft_v0.md"
FIG_LEGENDS = ROOT / "docs" / "figure_legends.md"
OUT = ROOT / "docs" / "manuscript_submission_ready.docx"

FIGURES = [
    ("Figure 1", ROOT / "results" / "figures_pub" / "Figure1_complete.png"),
    ("Figure 2", ROOT / "results" / "scrna" / "fig_umap_celltype_v2.png"),
    ("Figure 3", ROOT / "results" / "scrna" / "fig_GSE206647_grade_program.png"),
    ("Figure 4", ROOT / "results" / "figures_pub" / "fig_classical_drivers.png"),
    ("Figure 5", ROOT / "results" / "figures_pub" / "fig_geneformer_embedding.png"),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "D9E1EA")


def clean_inline(text):
    text = text.replace("**", "")
    text = text.replace("*", "")
    text = text.replace("`", "")
    text = re.sub(r"<[^>]+>", "", text)
    return text.strip()


def add_inline_runs(paragraph, text):
    text = text.replace("`", "")
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def enable_line_numbers(section):
    sect_pr = section._sectPr
    ln_num = sect_pr.find(qn("w:lnNumType"))
    if ln_num is None:
        ln_num = OxmlElement("w:lnNumType")
        sect_pr.append(ln_num)
    ln_num.set(qn("w:countBy"), "1")
    ln_num.set(qn("w:start"), "1")
    ln_num.set(qn("w:restart"), "newPage")


def add_page_number(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if paragraph.runs:
        return
    paragraph.add_run("Page ")
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])


def add_paragraph(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    add_inline_runs(p, text)
    return p


def set_portrait(section):
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    enable_line_numbers(section)
    add_page_number(section)


def set_landscape(section):
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    enable_line_numbers(section)
    add_page_number(section)


def add_markdown_table(doc, rows):
    parsed = []
    for row in rows:
        cells = [clean_inline(c) for c in row.strip().strip("|").split("|")]
        parsed.append(cells)
    if len(parsed) < 2:
        return
    header = parsed[0]
    body = [r for r in parsed[2:] if len(r) == len(header)]
    is_supp_cohort_table = header[:3] == ["Cohort", "Platform", "n (composition)"]
    if is_supp_cohort_table:
        set_landscape(doc.add_section(WD_SECTION.NEW_PAGE))
    table = doc.add_table(rows=1, cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = header[i]
        set_cell_shading(cell, "F2F4F7")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(7 if is_supp_cohort_table else 8)
    for row in body:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(7 if is_supp_cohort_table else 8)
    doc.add_paragraph()
    if is_supp_cohort_table:
        set_portrait(doc.add_section(WD_SECTION.NEW_PAGE))


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    enable_line_numbers(section)
    add_page_number(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 2.0

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_main_markdown(doc):
    lines = MANUSCRIPT.read_text().splitlines()
    table_buf = []
    skip_note = False
    for line in lines:
        if line.startswith("*Working draft"):
            continue
        if line.strip() == "---":
            continue
        if line.startswith("|"):
            table_buf.append(line)
            continue
        if table_buf:
            add_markdown_table(doc, table_buf)
            table_buf = []
        if not line.strip():
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(clean_inline(line[2:]))
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0, 0, 0)
        elif line.startswith("## "):
            doc.add_heading(clean_inline(line[3:]), level=1)
        elif line.startswith("### "):
            doc.add_heading(clean_inline(line[4:]), level=2)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, line[2:])
        else:
            add_paragraph(doc, line)
    if table_buf:
        add_markdown_table(doc, table_buf)


def add_figures(doc):
    doc.add_page_break()
    doc.add_heading("Figures", level=1)
    legends = FIG_LEGENDS.read_text().splitlines()
    legend_blocks = []
    current = []
    for line in legends:
        if line.startswith("**Figure ") and current:
            legend_blocks.append(current)
            current = [line]
        elif line.strip():
            current.append(line)
    if current:
        legend_blocks.append(current)

    legend_by_fig = {}
    for block in legend_blocks:
        title = clean_inline(block[0])
        key = title.split(".", 1)[0]
        legend_by_fig[key] = " ".join(clean_inline(x) for x in block)

    for label, path in FIGURES:
        if not path.exists():
            add_paragraph(doc, f"{label}: missing figure file {path}")
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(label).bold = True
        doc.add_picture(str(path), width=Inches(6.3))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_paragraph(doc, legend_by_fig.get(label, ""), style=None)
        doc.add_paragraph()


def main():
    doc = Document()
    setup_styles(doc)
    add_main_markdown(doc)
    add_figures(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
