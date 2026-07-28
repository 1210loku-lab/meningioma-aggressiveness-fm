#!/usr/bin/env python3
"""Append the author-supplied final JPEG figures to the submission manuscript.

This script treats ``figure_jpg/`` and ``figures_pdf_editable/`` as read-only
assets.  It writes only the combined author-review DOCX in ``manuscript/``.
"""

from pathlib import Path
from tempfile import TemporaryDirectory

from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
PACKAGE = ROOT / "submission" / "Scientific_Reports_20260715"
BASE = PACKAGE / "manuscript" / "Meningioma_Scientific_Reports_submission.docx"
FIGURE_DIR = PACKAGE / "figure_jpg"
OUTPUT = PACKAGE / "manuscript" / "Meningioma_Scientific_Reports_submission_with_all_figures.docx"

FIGURES = [
    ("Figure 1", "Fig1.jpg"),
    ("Figure 2", "Fig2.jpg"),
    ("Figure 3", "Fig3.jpg"),
    ("Figure 4", "Fig4.jpg"),
    ("Supplementary Figure S1", "FigS1.jpg"),
    ("Supplementary Figure S2", "FigS2.jpg"),
    ("Supplementary Figure S3", "FigS3.jpg"),
    ("Supplementary Figure S4", "FigS4.jpg"),
]


def add_page_number(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if footer.text:
        return
    footer.add_run("Page ")
    run = footer.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def main():
    if not BASE.exists():
        raise FileNotFoundError(BASE)
    missing = [name for _, name in FIGURES if not (FIGURE_DIR / name).exists()]
    if missing:
        raise FileNotFoundError(", ".join(str(FIGURE_DIR / name) for name in missing))

    document = Document(BASE)
    # The source manuscript carries an even-page header layout.  Disable it
    # only in this combined author-review copy: on landscape figure pages it
    # otherwise collides with the body label on every even page.
    document.settings.odd_and_even_pages_header_footer = False
    # python-docx cannot parse these valid, author-exported JPEG headers.  The
    # source assets remain untouched; temporary lossless PNG pixel copies are
    # used solely as Word-compatible embedding payloads.
    with TemporaryDirectory(prefix="meningioma_figures_") as temp_dir:
        temporary_dir = Path(temp_dir)
        section = document.add_section(WD_SECTION.NEW_PAGE)
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.top_margin = section.bottom_margin = Inches(0.45)
        section.left_margin = section.right_margin = Inches(0.5)
        section.header_distance = Inches(0.2)
        section.footer_distance = Inches(0.2)
        section.different_first_page_header_footer = False
        line_numbers = section._sectPr.find(qn("w:lnNumType"))
        if line_numbers is None:
            line_numbers = OxmlElement("w:lnNumType")
            section._sectPr.append(line_numbers)
        line_numbers.set(qn("w:countBy"), "0")
        add_page_number(section)

        for index, (label, filename) in enumerate(FIGURES):
            image_path = FIGURE_DIR / filename
            with Image.open(image_path) as image:
                ratio = image.width / image.height
                word_image = temporary_dir / f"{image_path.stem}.png"
                image.convert("RGB").save(word_image, "PNG")
            # Leave enough room for the figure label on every landscape page.
            max_width, max_height = 10.0, 5.70
            width = min(max_width, max_height * ratio)
            height = width / ratio
            if height > max_height:
                height = max_height
                width = height * ratio

            figure = document.add_paragraph()
            figure.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Retain the canonical Figure 1–5/S1–S2 order; the matching
            # legends remain in the manuscript.  Omit a separate body title
            # because the inherited even-page header collides with it in
            # Word/LibreOffice mixed-orientation rendering.
            figure.paragraph_format.page_break_before = index > 0
            figure.paragraph_format.space_after = Pt(0)
            figure.add_run().add_picture(
                str(word_image), width=Inches(width), height=Inches(height)
            )

        document.core_properties.author = "Mingyang Li, Qingyi Huo, Hui Wang and co-authors"
        document.core_properties.last_modified_by = "The authors"
        document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
